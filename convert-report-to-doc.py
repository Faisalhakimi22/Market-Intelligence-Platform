#!/usr/bin/env python3

'''
Simple script to convert the REPORT.md to REPORT.doc
Requires python-docx library: pip install python-docx markdown
'''

import os
import sys
from docx import Document
import markdown
import html2text

def main():
    try:
        # Check if python-docx is installed
        import docx
    except ImportError:
        print("The 'python-docx' library is required. Install it using 'pip install python-docx markdown'")
        sys.exit(1)
    
    # Check if the REPORT.md file exists
    if not os.path.exists('REPORT.md'):
        print("Error: REPORT.md file not found in the current directory.")
        sys.exit(1)
        
    try:
        # Read the markdown file
        with open('REPORT.md', 'r', encoding='utf-8') as md_file:
            md_content = md_file.read()
            
        # Convert markdown to HTML
        html_content = markdown.markdown(md_content)
        
        # Create a new Document object
        doc = Document()
        
        # Process the markdown content section by section
        sections = md_content.split('\n## ')
        
        # Process the title (first section)
        title = sections[0].strip().replace('# ', '')
        doc.add_heading(title, 0)
        
        # Process the remaining sections
        for i in range(1, len(sections)):
            section = sections[i]
            section_lines = section.strip().split('\n')
            section_title = section_lines[0]
            section_content = '\n'.join(section_lines[1:]).strip()
            
            # Add section heading
            doc.add_heading(section_title, 1)
            
            # Process subsections
            subsections = section_content.split('\n### ')
            
            # Add the main section text if there is any before the first subsection
            if not subsections[0].startswith('###'):
                paragraphs = subsections[0].strip().split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            # Process each subsection
            for j in range(1, len(subsections)):
                subsection = subsections[j]
                subsection_lines = subsection.strip().split('\n')
                subsection_title = subsection_lines[0]
                subsection_content = '\n'.join(subsection_lines[1:]).strip()
                
                # Add subsection heading
                doc.add_heading(subsection_title, 2)
                
                # Add subsection content
                if subsection_content:
                    paragraphs = subsection_content.split('\n\n')
                    for para in paragraphs:
                        if para.strip():
                            # Check if it's a list item
                            if para.strip().startswith('- '):
                                for list_item in para.strip().split('\n- '):
                                    if list_item.strip():
                                        if list_item.startswith('- '):
                                            list_item = list_item[2:]
                                        doc.add_paragraph(list_item.strip(), style='List Bullet')
                            else:
                                doc.add_paragraph(para.strip())
        
        # Save the document
        doc.save('REPORT.docx')
        print("Successfully converted REPORT.md to REPORT.docx")
        
    except Exception as e:
        print(f"Error converting the report: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()
